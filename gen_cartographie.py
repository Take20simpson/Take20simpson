#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# Colors
NAVY      = RGBColor(24, 54, 99)    # #183663 - main title
BLUE1     = RGBColor(0, 70, 127)    # #00467F - section headers
BLUE2     = RGBColor(46, 116, 181)  # #2E74B5 - bloc headers
BLUE3     = RGBColor(68, 114, 196)  # #4472C4 - sub-headers / à venir
GRAY      = RGBColor(89, 89, 89)
RED_SOFT  = RGBColor(192, 0, 0)
GREEN_SOFT= RGBColor(0, 112, 0)
ORANGE    = RGBColor(197, 90, 17)

# ─── helpers ────────────────────────────────────────────────

def sp(p, before=0, after=5):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def title(doc, text, size=22, color=NAVY, center=True, before=8, after=4):
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.color.rgb = color; r.font.size = Pt(size)
    sp(p, before, after)
    return p

def section(doc, text, size=15, color=BLUE1, before=18, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.color.rgb = color; r.font.size = Pt(size)
    sp(p, before, after)
    return p

def bloc(doc, text, size=12.5, color=BLUE2, before=12, after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.color.rgb = color; r.font.size = Pt(size)
    sp(p, before, after)
    return p

def sub(doc, text, size=11, color=BLUE3, before=8, after=3):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.color.rgb = color; r.font.size = Pt(size)
    sp(p, before, after)
    return p

def bullet(doc, runs, level=1, before=0, after=3):
    """runs = list of (text, bold, italic, underline, color) or just str"""
    p = doc.add_paragraph()
    indent = Cm(level * 0.65)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.42)
    sp(p, before, after)
    br = p.add_run("•  ")
    br.font.size = Pt(11); br.font.color.rgb = BLUE2
    if isinstance(runs, str):
        r = p.add_run(runs); r.font.size = Pt(11)
    elif isinstance(runs, list):
        for item in runs:
            if isinstance(item, str):
                r = p.add_run(item); r.font.size = Pt(11)
            else:
                txt, bold, italic, underline, color = item
                r = p.add_run(txt)
                r.bold = bold; r.italic = italic; r.underline = underline
                r.font.size = Pt(11)
                if color: r.font.color.rgb = color
    return p

def b(text): return (text, True, False, False, None)
def i(text): return (text, False, True, False, GRAY)
def u(text): return (text, True, False, True, None)
def bi(text): return (text, True, True, False, None)
def n(text): return (text, False, False, False, None)
def c(text, color): return (text, False, False, False, color)

def note(doc, text, level=1, before=2, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.65 + 0.42)
    sp(p, before, after)
    r = p.add_run(text); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
    return p

def rule(doc, text, level=1, before=4, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.65)
    sp(p, before, after)
    r = p.add_run("⚠  " + text)
    r.bold = True; r.underline = True; r.font.size = Pt(10.5)
    return p

def separator(doc, before=8, after=8):
    p = doc.add_paragraph()
    sp(p, before, after)
    r = p.add_run("─" * 85)
    r.font.color.rgb = RGBColor(189, 189, 189); r.font.size = Pt(8)
    return p

def normal(doc, runs, left=0.65, before=0, after=5):
    p = doc.add_paragraph()
    if left: p.paragraph_format.left_indent = Cm(left)
    sp(p, before, after)
    if isinstance(runs, str):
        r = p.add_run(runs); r.font.size = Pt(11)
    else:
        for item in runs:
            if isinstance(item, str):
                r = p.add_run(item); r.font.size = Pt(11)
            else:
                txt, bold, italic, underline, color = item
                r = p.add_run(txt)
                r.bold = bold; r.italic = italic; r.underline = underline
                r.font.size = Pt(11)
                if color: r.font.color.rgb = color
    return p

def mono(doc, text, before=0, after=0):
    p = doc.add_paragraph()
    sp(p, before, after)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(8.5)
    return p

def step_box(doc, text, color=BLUE1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, 3, 3)
    r = p.add_run(f"【  {text}  】")
    r.bold = True; r.font.color.rgb = color; r.font.size = Pt(11)
    return p

def arrow(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, 1, 1)
    r = p.add_run("↓")
    r.font.color.rgb = GRAY; r.font.size = Pt(13)
    return p

# ════════════════════════════════════════════════════════════
#  DOCUMENT
# ════════════════════════════════════════════════════════════

# ── TITRE ──────────────────────────────────────────────────
title(doc, "CARTOGRAPHIE COMPLÈTE", size=24, before=4, after=2)
title(doc, "SETTING AGENT IA  —  IA DU LINK", size=16, color=BLUE1, before=0, after=3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, 0, 14)
r = p.add_run("Cartographie stratégique — ce qui va où, pourquoi, et comment")
r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(11)

separator(doc, 4, 16)

# ── SECTION 0 : L'ENJEU CENTRAL ────────────────────────────
section(doc, "▌  L'ENJEU CENTRAL : FEELING → FRAMEWORK")

normal(doc, [
    n("Matthias possède un "),
    bi("feeling naturel"),
    n(" pour le setting LinkedIn. Il sent instinctivement quand creuser, quand reculer, comment formuler pour que ça atterrisse. Ce que l'IA doit accomplir : "),
    b("cartographier ce feeling"),
    n(" en règles claires, en signaux détectables, en séquences reproductibles."),
], left=0.65, after=8)

bullet(doc, "Ce n'est pas un problème de technique — c'est un problème de transcription")
bullet(doc, "Le tronc central n'est pas un script. C'est une architecture de décision.")
bullet(doc, "L'objectif : que l'IA réplique le feeling, pas qu'elle simule un vendeur.")

separator(doc)

# ── SECTION 1 : ARCHITECTURE ───────────────────────────────
section(doc, "1.  ARCHITECTURE GLOBALE — 4 NIVEAUX")

bloc(doc, "NIVEAU 1 — TRONC CENTRAL")
bullet(doc, "Universel, niche-agnostic — s'applique quelle que soit la niche")
bullet(doc, "8 blocs structurants (voir Section 2)")
bullet(doc, "C'est la fondation — tout le reste s'y greffe")

bloc(doc, "NIVEAU 2 — MATÉRIAU D'ENTRAÎNEMENT")
bullet(doc, "15+ conversations réelles de Matthias analysées")
bullet(doc, [n("Format : "), b("contexte → signal détecté → action choisie → raisonnement")])
bullet(doc, "Alimente le Bloc 3 (système de signaux) avec des cas réels")
bullet(doc, "Ce n'est pas du contenu « à lire » — c'est de la data d'apprentissage")

bloc(doc, "NIVEAU 3 — MODULES NICHE")
bullet(doc, "Superpositions spécifiques greffées sur le tronc central")
bullet(doc, "Ex : entonnoir corporate vs solopreneur")
bullet(doc, "Ex : méthodes de qualification prix par profil")
bullet(doc, "Ex : pain points par secteur d'activité")
note(doc, "→ À construire APRÈS le tronc central et l'analyse des conversations")

bloc(doc, "NIVEAU 4 — CONFIGURATION UTILISATEUR")
bullet(doc, "Ce que chaque utilisateur de la plateforme configure lui-même")
bullet(doc, "Histoires personnelles, offre, ton propre, seuils")
bullet(doc, "Tics de langage, expressions récurrentes")
note(doc, "→ Pas dans le tronc central — l'utilisateur peut le modifier seul")

separator(doc)

# ── SECTION 2 : LES 8 BLOCS ────────────────────────────────
section(doc, "2.  LE TRONC CENTRAL — LES 8 BLOCS")

# ─ BLOC 1 ─
bloc(doc, "BLOC 1 — Vision macro des phases")

bullet(doc, "3 grandes phases du setting :")
bullet(doc, [b("Phase 1 : "), n("entrer dans la conversation (CAS A ou CAS B)")], level=2)
bullet(doc, [b("Phase 2 : "), n("conscientiser le problème (questions de fond, signaux)")], level=2)
bullet(doc, [b("Phase 3 : "), n("qualifier et proposer l'action finale")], level=2)

rule(doc, "On ne passe d'une phase à l'autre qu'avec des critères de transition clairs — jamais par automatisme")

bullet(doc, [b("CAS A : "), n("contact après commentaire sur un post de l'utilisateur — le prospect vient à lui")])
bullet(doc, [b("CAS B : "), n("contact en outbound — l'utilisateur va vers le prospect")])
bullet(doc, [b("CAS C : "), n("le prospect envoie une demande de connexion — initiative entrante du prospect")])
bullet(doc, [b("CAS D : "), n("visiteur de profil — le prospect a visité le profil de l'utilisateur sans envoyer de message")])
note(doc, "→ Les 4 cas ne suivent pas le même flux d'ouverture, mais convergent à partir de la Phase 2")

# ─ BLOC 2 ─
bloc(doc, "BLOC 2 — Profils personnalité")

bullet(doc, "Deux grands profils (adapté DISC au contexte LinkedIn) :")
bullet(doc, [b("Profil transactionnel (rouge) : "), n("direct, efficace, peu bavard, peu expressif")], level=2)
bullet(doc, [b("Profil conversationnel : "), n("ouvert, expressif, aime échanger, partage facilement")], level=2)

bullet(doc, "Comment détecter :")
bullet(doc, "Comportement en commentaires (longueur, registre)", level=2)
bullet(doc, "Ton des posts (distant/factuel vs chaleureux/narratif)", level=2)
bullet(doc, "Longueur et profondeur des réponses en DM", level=2)

bullet(doc, "Ce que ça détermine :")
bullet(doc, "Le registre adopté (direct vs chaleureux)", level=2)
bullet(doc, "Le rythme de la conversation (aller vite vs prendre le temps)", level=2)
bullet(doc, "La profondeur de questions acceptable sans brusquer", level=2)

# ─ BLOC 3 ─
bloc(doc, "BLOC 3 — Système de signaux")

normal(doc, [b("Principe : "), n("IF [signal détecté]  →  THEN [action + raisonnement]")], after=6)

bullet(doc, "Exemples de signaux et leurs actions associées :")
bullet(doc, [b("Signal d'engagement fort "), n("→ approfondir, ne pas changer de sujet")], level=2)
bullet(doc, [b("Signal d'évasion "), n("→ changer d'angle doucement, ne pas insister sur le même point")], level=2)
bullet(doc, [b("Signal d'histoire personnelle partagée "), n("→ réciprocité contextuelle (partager une histoire en miroir)")], level=2)
bullet(doc, [b("Signal de désengagement "), n("→ ne pas relancer de force, laisser respirer")], level=2)
bullet(doc, [b("Signal de profil silencieux "), n("→ proposer alternatives A/B (à documenter lors de l'analyse des conversations)")], level=2)

rule(doc, "Les histoires personnelles sont déclenchées par signal — jamais par fréquence mécanique (ex : « 1 fois sur 4 » = interdit)")
rule(doc, "Garde-fou de fréquence : même quand un signal est détecté, l'IA ne partage pas une histoire personnelle plus d'une fois sur 3 à 4 messages — détecter un bon pattern ne suffit pas à le systématiser")
note(doc, "→ Double logique : (1) c'est le signal qui déclenche, pas une règle de comptage — ET (2) même avec un signal, on ne le fait pas à chaque fois")
note(doc, "→ Le contenu précis de ce bloc sera enrichi par l'analyse des 15 conversations réelles")

# ─ BLOC 4A ─
bloc(doc, "BLOC 4A — Ultra-personnalisation par insight profond")

normal(doc, [
    b("Principe : "),
    n("synthétiser une lecture profonde du profil en "),
    u("UN insight"),
    n(" qui montre qu'on a vraiment compris la personne — pas juste lu son titre LinkedIn"),
], after=6)

bullet(doc, "Ce qu'on lit — tout ce qui est disponible en texte :")
bullet(doc, "Posts, section info, commentaires, stratégie visible sur le profil", level=2)
note(doc, "→ L'agent ne peut pas lire les vidéos — il s'appuie sur le maximum d'informations textuelles disponibles", level=2)
bullet(doc, "Pas juste le headline — la cohérence globale de la présence LinkedIn", level=2)

bullet(doc, [c("✗  ", RED_SOFT), b("Ce qu'on NE fait PAS : "), i('"J\'ai vu sur ton profil que tu faisais X, Y, Z"'), n(" → mécanique, éclaté au sol")])
bullet(doc, [c("✓  ", GREEN_SOFT), b("Ce qu'on fait : "), n("arriver avec "), b("LA synthèse"), n(' qui fait dire "putain, ce mec a vraiment regardé mes posts, compris ma vision et ma stratégie"')])

bullet(doc, [b("Test de validation : "), n("la personne penserait-elle « il a compris ma philosophie » — pas juste « il a lu mon profil » ?")])

bullet(doc, "Moments d'application :")
bullet(doc, "Ouverture CAS B (2ème message après connexion)", level=2)
bullet(doc, "Question de transition entre phases", level=2)
bullet(doc, "Parfois dans les questions de fond", level=2)

# ─ BLOC 4 ─
bloc(doc, "BLOC 4 — Règles sur les questions")

bullet(doc, [b("Défaut absolu : "), n("question ouverte — toujours")])
bullet(doc, [b("Question fermée : "), n("seulement avec critères contextuels très précis (à définir lors de l'analyse des conversations)")])
rule(doc, "Une seule question par message — jamais deux en même temps")
bullet(doc, [b("Profondeur > largeur : "), n("rester sur un sujet jusqu'à avoir quelque chose d'émotionnel avant de passer au suivant")])
bullet(doc, "Les questions servent à faire conscientiser — pas à collecter des informations")

# ─ BLOC 5 ─
bloc(doc, "BLOC 5 — Standards qualité")

bullet(doc, "Ton authentique, non-corporate, non-template")
bullet(doc, "Ne jamais signaler mécaniquement qu'on a lu le profil")
bullet(doc, "Variation maximale des formulations (pas de patterns répétitifs détectables)")
bullet(doc, "Équilibre : chaleureux mais jamais familier — direct mais jamais sec")

# ─ BLOC 6 ─
bloc(doc, "BLOC 6 — Qualification")

normal(doc, [
    b("Objectif : confirmer "),
    u("besoin réel"),
    b(" + confirmer "),
    u("solvabilité"),
    b(" — deux axes "),
    u("indépendants"),
    b(" à valider séparément"),
], after=7)

bullet(doc, [u("Besoin réel / Problème existant"), n(" = la personne a un problème identifié et reconnu")])
bullet(doc, [i('"Mes clients sont tous en delivery, je n\'ai pas besoin là"'), n(" → besoin absent ou non perçu même si solvabilité OK")], level=2)

bullet(doc, [u("Solvabilité"), n(" = la personne peut investir")])
bullet(doc, [i('"Pas de budget"'), n(" ≠ pas de solvabilité → souvent une résistance à verbaliser, pas une réalité financière")], level=2)
note(doc, "→ Comment distinguer les deux selon la niche : dans les modules niche (corporate vs solopreneur)", level=2)

bullet(doc, [b("Temporalité du besoin : "), n("quelqu'un peut avoir un besoin réel et les moyens, mais pas le bon timing — à vérifier discrètement")])
note(doc, "→ À intégrer comme vérification discrète dans la qualification, sans en faire une grande règle formelle")

bullet(doc, "Méthodes universelles présentes dans le tronc central :")
bullet(doc, "Storytelling approach (qualifier indirectement par projection)", level=2)
bullet(doc, "Calendar filter (filtrer par disponibilité)", level=2)
bullet(doc, "Smooth framing (recadrer naturellement la conversation vers la qualification)", level=2)

rule(doc, "Méthodes spécifiques par profil (corporate = question directe possible / solopreneur = indirect) → dans les modules niche, pas ici")
bullet(doc, [b("Lead scoring : "), n("critère minimum à atteindre avant de proposer l'action finale")])

# ─ BLOC 7 ─
bloc(doc, "BLOC 7 — Signaux de disqualification")

normal(doc, [
    b("Principe clé : "),
    n("l'IA détecte et signale — "),
    u("la décision finale reste chez l'utilisateur"),
], after=6)

bullet(doc, "Signaux types :")
bullet(doc, "Même évasion répétée 4 fois → signal d'alerte", level=2)
bullet(doc, [i('"Je me débrouille seul avec Claude"'), n(" → signal possible de non-alignement")], level=2)
bullet(doc, "Manque d'honnêteté flagrant dans les réponses", level=2)
bullet(doc, "Résistance systématique à toute question, quel que soit l'angle", level=2)

rule(doc, "Certains profils difficiles valent le coup de pousser — d'autres non. C'est le jugement de l'utilisateur, pas de l'IA.")
note(doc, "→ Contenu précis des signaux à enrichir lors de l'analyse des conversations réelles")

separator(doc)

# ── SECTION 3 : SÉQUENCES CLÉS ─────────────────────────────
section(doc, "3.  SÉQUENCES CLÉS À RETENIR")

bloc(doc, "Séquence centrale : Conscientisation → Permission → Mini-insight")

step_box(doc, "Questions de conscientisation", BLUE1)
arrow(doc)
step_box(doc, '"Est-ce que tu m\'autorises à te dire quelque chose ?"', BLUE2)
arrow(doc)
step_box(doc, 'Mini-insight — "Tu n\'as pas vraiment un problème de X, tu as un problème de Y"', BLUE1)

p = doc.add_paragraph(); sp(p, 8, 4)

bullet(doc, "La personne a verbalisé et conscientisé son problème via les questions de fond")
bullet(doc, "La demande de permission crée un micro-engagement psychologique")
note(doc, "La personne dit oui → elle s'ouvre → le reframe atterrit infiniment mieux qu'en le balançant directement")
bullet(doc, [b("Mini-insight : "), n("nommer le vrai problème derrière le problème apparent")])

sub(doc, "À venir : Entonnoir corporate vs solopreneur", color=BLUE3, before=14)
bullet(doc, [b("Corporate : "), n("qualifier tôt + personnalisation profonde + résultats concrets en référence")])
bullet(doc, [b("Solopreneur : "), n("discuter d'abord, qualifier en fin de conversation")])
note(doc, "→ À affiner dans les modules niche — pas dans le tronc central")

separator(doc)

# ── SECTION 4 : CE QUI N'EST PAS DANS LE TRONC CENTRAL ────
section(doc, "4.  CE QUI N'EST PAS DANS LE TRONC CENTRAL")

bloc(doc, "Dans les modules niche :")
bullet(doc, "Entonnoir corporate vs solopreneur (direction et séquence selon profil)")
bullet(doc, "Méthodes de qualification prix par profil")
bullet(doc, "Pain points spécifiques par secteur d'activité")
bullet(doc, "Critères de scoring initial du profil → quel entonnoir activer")

bloc(doc, "Dans la configuration utilisateur :")
bullet(doc, "Histoires personnelles de l'utilisateur de la plateforme")
bullet(doc, "Offre et positionnement")
bullet(doc, "Tics de langage, expressions récurrentes, ton propre")
bullet(doc, "Seuils de qualification personnels")

bloc(doc, "Pour plus tard — lors de l'analyse des conversations :")
bullet(doc, "Alternatives A/B pour profils silencieux → à documenter avec des cas réels")
bullet(doc, "Contenu précis des signaux (Bloc 3 et Bloc 7) → à enrichir avec les patterns détectés")

separator(doc)

# ── SECTION 5 : ORDRE D'EXÉCUTION ──────────────────────────
section(doc, "5.  ORDRE D'EXÉCUTION")

steps = [
    ("1", "Écrire le tronc central complet",
     "TRONC_CENTRAL_SETTING_AGENT_IA.md — les 8 blocs avec toutes les règles"),
    ("2", "Analyser 15+ conversations réelles",
     "Extraire les signaux → enrichir Bloc 3 (système de signaux) et Bloc 7 (disqualification)"),
    ("3", "Construire les modules niche",
     "Corporate / solopreneur — entonnoirs et méthodes de qualification"),
    ("4", "Configuration utilisateur",
     "Framework personnalisé pour chaque utilisateur de la plateforme"),
]

for num, t, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.65)
    sp(p, 5, 1)
    r1 = p.add_run(f"  {num}.  ")
    r1.bold = True; r1.font.color.rgb = BLUE1; r1.font.size = Pt(13)
    r2 = p.add_run(t)
    r2.bold = True; r2.font.size = Pt(12)
    note(doc, desc, level=2, before=1, after=6)

separator(doc)

# ── SECTION 6 : SCHÉMA ─────────────────────────────────────
section(doc, "6.  SCHÉMA D'ARCHITECTURE")

schema = [
    "┌──────────────────────────────────────────────────────────────┐",
    "│                      TRONC CENTRAL                          │",
    "│  ──────────────────────────────────────────────────────────  │",
    "│  Bloc 1  — Vision macro des phases                          │",
    "│  Bloc 2  — Profils personnalité (DISC)                      │",
    "│  Bloc 3  — Système de signaux (IF signal → THEN action)     │",
    "│  Bloc 4A — Ultra-personnalisation par insight profond       │",
    "│  Bloc 4  — Règles sur les questions                         │",
    "│  Bloc 5  — Standards qualité                                │",
    "│  Bloc 6  — Qualification (besoin réel + solvabilité)        │",
    "│  Bloc 7  — Signaux de disqualification                      │",
    "└────────────────────────────┬─────────────────────────────────┘",
    "                             │  alimenté par",
    "┌────────────────────────────▼─────────────────────────────────┐",
    "│                 MATÉRIAU D'ENTRAÎNEMENT                     │",
    "│  15+ conversations réelles analysées                        │",
    "│  Format : signal → action → raisonnement                    │",
    "└────────────────────────────┬─────────────────────────────────┘",
    "                             │  enrichi par / enrichit",
    "┌────────────────────────────▼─────────────────────────────────┐",
    "│                     MODULES NICHE                           │",
    "│  Corporate  :  entonnoir qualification précoce              │",
    "│  Solopreneur:  discussion d'abord, qualifier à la fin       │",
    "│  Qualification prix par profil                              │",
    "│  Pain points par métier / secteur                           │",
    "└──────────────────────────────────────────────────────────────┘",
    "",
    "┌──────────────────────────────────────────────────────────────┐",
    "│               CONFIGURATION UTILISATEUR                     │",
    "│  Histoires perso, offre, ton propre, seuils                 │",
    "│  Tics de langage, expressions récurrentes                   │",
    "└──────────────────────────────────────────────────────────────┘",
]

for line in schema:
    mono(doc, line)

# save
doc.save('/home/user/Take20simpson/cartographie_setting_agent.docx')
print("Saved.")
