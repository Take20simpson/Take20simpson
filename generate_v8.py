#!/usr/bin/env python3
"""
Génération de TRONC_CENTRAL_Yadulink_v8.docx à partir de v7.docx
"""

import shutil
import zipfile
import os
import re
from copy import deepcopy
from lxml import etree

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = '/root/.claude/uploads/07d5639f-0a6f-4f92-abfa-161acae1e8c7/32b9251e-TRONC_CENTRAL_Yadulink_v7.docx'
DST = '/home/user/Take20simpson/TRONC_CENTRAL_Yadulink_v8.docx'
TMP = '/tmp/v8_work.docx'

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 0 : copie du fichier source
# ──────────────────────────────────────────────────────────────────────────────
shutil.copy2(SRC, TMP)
doc = Document(TMP)

# ──────────────────────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def para_full_text(p):
    return ''.join(run.text for run in p.runs)

def set_run_bold(run, bold):
    run.bold = bold

def set_run_italic(run, italic):
    run.italic = italic

def set_run_underline(run, underline):
    run.underline = underline

def set_run_color(run, rgb_tuple):
    run.font.color.rgb = RGBColor(*rgb_tuple)

def set_run_size(run, pt):
    run.font.size = Pt(pt)

def apply_heading_style(para, level, color_rgb, size_pt, space_before_pt, space_after_pt):
    """Apply heading style formatting directly to paragraph XML."""
    style_name = f'Heading {level}'
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass

    pPr = para._p.get_or_add_pPr()
    # Space before/after
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(space_before_pt * 20)))
    spacing.set(qn('w:after'), str(int(space_after_pt * 20)))

    # Apply formatting to all runs
    for run in para.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
        run.font.size = Pt(size_pt)
        if level == 1:
            run.bold = True
        elif level == 2:
            run.bold = True
        elif level == 3:
            run.bold = True
            run.underline = True

def merge_runs_text(para):
    """Return the full text of a paragraph."""
    return para.text

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 1 : CORRECTIONS DE TEXTE
# ──────────────────────────────────────────────────────────────────────────────

text_replacements = [
    ("L'rondeur", "La rondeur"),
    ("l'rondeur", "la rondeur"),
    ("d'rondeur", "de rondeur"),
    ("cet rondeur", "cette rondeur"),
]

print("Étape 1 : corrections texte...")
for para in doc.paragraphs:
    for run in para.runs:
        for old, new in text_replacements:
            if old in run.text:
                run.text = run.text.replace(old, new)

# Rename ENROBAGE → RONDEUR in Bloc 3 vocabulary entry label
# Para 328 has "ENROBAGE\nTout élément..." - need to rename heading and Heading 2 of SECTION 2
for para in doc.paragraphs:
    txt = para.text
    # The vocab entry ENROBAGE (in vocab bloc 3, not elsewhere)
    # Check if it's a normal para with just ENROBAGE and a definition
    if txt.strip().startswith('ENROBAGE') and 'Tout élément conversationnel' in txt:
        for run in para.runs:
            if 'ENROBAGE' in run.text:
                run.text = run.text.replace('ENROBAGE', 'RONDEUR')
    # The Heading 3 "ENROBAGE" in vocabulary section
    if txt.strip() == 'ENROBAGE' and para.style.name in ('Heading 3', 'normal', 'Normal'):
        # Only rename in Bloc 3 vocabulary context
        for run in para.runs:
            if run.text.strip() == 'ENROBAGE':
                run.text = run.text.replace('ENROBAGE', 'RONDEUR')
    # Heading 2 "SECTION 2 , L'ENROBAGE" → "SECTION 2 , LA RONDEUR"
    if "SECTION 2 , L'ENROBAGE" in txt:
        for run in para.runs:
            run.text = run.text.replace("SECTION 2 , L'ENROBAGE", "SECTION 2 , LA RONDEUR")
    # "Deux formes principales :" line to remove — handled in step 2 (phantom removal)

print("  -> Corrections texte appliquées")

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 2 : SUPPRESSION DE LA SECTION 4.5 FANTÔME (Bloc 3, paras 485-539)
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 2 : suppression section 4.5 fantôme du Bloc 3...")

paragraphs = doc.paragraphs
# Find the phantom 4.5 in Bloc 3 (around para 485)
# It starts with '4.5 , Suivi des prospects' and ends before 'SECTION 5'
# We need to find it by content since indices may shift

phantom_start_idx = None
phantom_end_idx = None

for i, p in enumerate(paragraphs):
    txt = p.text
    if phantom_start_idx is None:
        # Find first occurrence that's NOT a Heading 3 (the heading version at ~1378 is OK to keep)
        if '4.5 , Suivi des prospects' in txt and p.style.name not in ('Heading 3',):
            # Check it's in Bloc 3 zone (before SECTION 5 which is at ~540)
            phantom_start_idx = i
    elif phantom_start_idx is not None and phantom_end_idx is None:
        if p.style.name == 'Heading 2' and 'SECTION 5' in txt:
            phantom_end_idx = i
            break

print(f"  Phantom 4.5: paras {phantom_start_idx} to {phantom_end_idx-1}")

if phantom_start_idx is not None and phantom_end_idx is not None:
    # Remove paragraphs from phantom_start_idx to phantom_end_idx-1 (inclusive)
    # We work with the XML elements
    body = doc.element.body
    # Get all paragraph elements in order
    all_paras = doc.paragraphs
    paras_to_remove = all_paras[phantom_start_idx:phantom_end_idx]
    for p in paras_to_remove:
        p._element.getparent().remove(p._element)
    print(f"  -> {phantom_end_idx - phantom_start_idx} paragraphes supprimés")

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 3 : DÉPLACEMENT VOCABULAIRE DU BLOC 2
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 3 : déplacement vocabulaire Bloc 2...")

# After step 2, indices have shifted. Re-read paragraphs.
paragraphs = doc.paragraphs

# Find VOCABULAIRE DU BLOC 2 heading paragraph
vocab_bloc2_start = None
vocab_bloc2_end = None
bloc2_title_idx = None

for i, p in enumerate(paragraphs):
    txt = p.text.strip()
    if 'VOCABULAIRE DU BLOC 2' in txt and vocab_bloc2_start is None:
        vocab_bloc2_start = i
    elif vocab_bloc2_start is not None and vocab_bloc2_end is None:
        if p.style.name == 'Heading 1' and 'BLOC 2' in txt:
            vocab_bloc2_end = i
            bloc2_title_idx = i

print(f"  Vocab Bloc2: {vocab_bloc2_start} to {vocab_bloc2_end-1}, BLOC2 title: {bloc2_title_idx}")

if vocab_bloc2_start is not None and vocab_bloc2_end is not None:
    # Extract vocab bloc 2 elements
    vocab_elements = [p._element for p in paragraphs[vocab_bloc2_start:vocab_bloc2_end]]
    # Clone them
    clones = [deepcopy(e) for e in vocab_elements]

    # Remove originals
    for e in vocab_elements:
        e.getparent().remove(e)

    # Re-read paragraphs to find BLOC 2 heading new position
    paragraphs = doc.paragraphs
    bloc2_para = None
    for p in paragraphs:
        if p.style.name == 'Heading 1' and 'BLOC 2 , PROFILS' in p.text:
            bloc2_para = p
            break

    if bloc2_para is not None:
        # Insert clones after BLOC 2 heading
        ref = bloc2_para._element
        parent = ref.getparent()
        idx_in_parent = list(parent).index(ref)
        for j, clone in enumerate(clones):
            parent.insert(idx_in_parent + 1 + j, clone)
        print("  -> Vocabulaire Bloc 2 déplacé après BLOC 2 heading")

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 4 : AJOUT EMOJI 🫩 À LA LISTE AUTORISÉE (Bloc 5)
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 4 : ajout emoji 🫩...")

paragraphs = doc.paragraphs
for p in paragraphs:
    if '😊 😄 😅 😂 🙂 😉 😬 🙈 :) ;) :D ahah/haha 👍 🙏 💪 ✅ 👏 🔥' in p.text:
        for run in p.runs:
            if '🔥' in run.text and '🫩' not in run.text:
                run.text = run.text + ' 🫩'
                print("  -> 🫩 ajouté")
                break

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 5 : STYLES DE TITRES
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 5 : application des styles de titres...")

H1_TITLES = {
    "TRONC CENTRAL : STRATÉGIE DE SETTING AGENT IA",
    "LEXIQUE , VOCABULAIRE CENTRAL",
    "BLOC 1 , VISION MACRO DES PHASES + 4 CAS D'ENTRÉE",
    "BLOC 2 , PROFILS PERSONNALITÉ",
    "BLOC 3 , SYSTÈME DE SIGNAUX",
    "BLOC 4A , ULTRA-PERSONNALISATION PAR INSIGHT PROFOND",
    "BLOC 4B , RÈGLES SUR LES QUESTIONS",
    "BLOC 5 , STANDARDS QUALITÉ",
    "BLOC 6 , QUALIFICATION",
    "BLOC 7 , SIGNAUX DE DISQUALIFICATION",
}

H2_TITLES = {
    "VOCABULAIRE DU BLOC",  # prefix match
    "PRINCIPE CENTRAL",
    "NUANCE CENTRALE , LES DEUX ENTONNOIRS SONT CONVERSATIONNELS",
    "RÈGLE CENTRALE , PHASES 1 ET 2 (ENTONNOIR CLASSIQUE)",
    "LES DEUX CHEMINS",
    "LES 3 PHASES",
    "CRITÈRES DE TRANSITION ENTRE PHASES",
    "LES 4 CAS D'ENTRÉE",
    "POINT DE CONVERGENCE",
    "LES DEUX COUCHES D'ANALYSE",
    "LES SIGNAUX DE DÉTECTION",
    "LES 6 ARCHÉTYPES",
    "CE QUE LE PROFIL DÉTERMINE",
    "LA RECLASSIFICATION EN COURS DE CONVERSATION",
    "LE FORMAT CENTRAL",
    "SECTION 1 , CLASSIFICATION DES RÉPONSES ET ACTIONS ASSOCIÉES",
    "SECTION 2 , L'ENROBAGE",
    "SECTION 2 , LA RONDEUR",
    "SECTION 3 , SIGNAUX D'ENGAGEMENT ÉMOTIONNEL",
    "SECTION 4 , SIGNAUX DE DOULEUR ÉMOTIONNELLE",
    "SECTION 5 , GESTION DES RÉPONSES MINIMISANTES",
    "SECTION 6 , LE DÉTOUR CONVERSATIONNEL",
    "SECTION 7 , EFFET MIROIR , AVERTISSEMENT",
    "SECTION 8 , PROFILS SILENCIEUX",
    "CE QU'ON LIT , SOURCES ET HIÉRARCHIE",
    "CE QU'ON NE FAIT PAS , CE QU'ON FAIT",
    "LES 3 NIVEAUX DE FORMULATION",
    "LONGUEUR",
    "TEST DE VALIDATION",
    "MOMENTS D'APPLICATION , RÈGLES PAR CAS",
    "SUJETS À ÉVITER DANS L'INSIGHT",
    "RÈGLE 1, QUESTIONS OUVERTES PAR DÉFAUT",
    "RÈGLE 2, QUESTIONS FERMÉES : USAGE STRICTEMENT LIMITÉ",
    "RÈGLE 3, UNE UNITÉ DE QUESTION PAR MESSAGE",
    "RÈGLE 4, LES DEUX TYPES DE QUESTIONS",
    "RÈGLE 5, PROFONDEUR > LARGEUR",
    "RÈGLE 6, LA QUESTION FERMÉE AU MAUVAIS MOMENT",
    "L'ORDRE NATUREL DES QUESTIONS",
    "SECTION 1 , LE REBOND D'ÉNERGIE",
    "SECTION 2 , VARIATION DES FORMULATIONS",
    "SECTION 3 , RONDEUR BIEN PLACÉE VS MAL PLACÉE",
    "SECTION 4 , EXPRESSIONS ET PATTERNS À BANNIR",
    "SECTION 5 , PONCTUATION, EMOJIS, IMPERFECTIONS",
    "SECTION 6 , TEST DE VALIDATION",
    "SECTION 1 , BESOIN RÉEL",
    "SECTION 2 , SOLVABILITÉ",
    "SECTION 3 , LEAD SCORING",
    "LE DÉCLENCHEUR DE PROPOSITION",
    "SECTION 4 , TEMPORALITÉ ET GESTION DES OBJECTIONS",
    "SECTION 5 , LA STORYTELLING APPROACH",
    "SECTION 6 , POINTS D'ATTENTION TRANSVERSAUX",
    "SECTION 1 , SIGNAUX DE DISQUALIFICATION RÉELS",
    "SECTION 2 , SIGNAUX D'ALERTE → ESPACE DE VÉRIFICATION UTILISATEUR",
    "SECTION 3 , CE QUI N'EST PAS UNE DISQUALIFICATION",
    "SECTION 4 , POINTS D'ATTENTION TRANSVERSAUX",
    "SÉQUENCE EN 3 ÉTAPES",
}

H3_TITLES = {
    "ENTONNOIR CLASSIQUE", "ENTONNOIR INVERSÉ", "ACTION FINALE", "CONSCIENTISATION",
    "PERMISSION", "MINI-INSIGHT",
    "PHASE 1 , Entrée dans la conversation", "PHASE 2 , Transition / Focus",
    "PHASE 3 , Profondeur / Qualification",
    "Phase 1 → Phase 2 (entonnoir classique uniquement)",
    "Phase 1 → Phase 3 (entonnoir inversé uniquement)",
    "Phase 2 → Phase 3 (entonnoir classique)",
    "CAS A , Après échanges en commentaire sous un post du prospect",
    "CAS B , Outbound", "CAS C , Demande de connexion entrante", "CAS D , Visiteur de profil",
    "PROFIL TRANSACTIONNEL", "PROFIL CONVERSATIONNEL", "PROSPECT AVERTI", "RECLASSIFICATION",
    "Signaux du profil transactionnel", "Signaux du profil conversationnel",
    "ARCHÉTYPE 1 , La freelance débutante", "ARCHÉTYPE 2 , Le/la freelance expérimenté(e)",
    "ARCHÉTYPE 3 , Le coach / consultant(e)", "ARCHÉTYPE 4 , Le corporate / décideur",
    "ARCHÉTYPE 5 , Le fondateur / builder",
    "ARCHÉTYPE 6 , Le prospect averti  ⚠ COUCHE TRANSVERSALE",
    "Principe", "Le seul switch qui compte : vers la qualification directe",
    "ENROBAGE", "RONDEUR", "SIGNAL D'ENGAGEMENT ÉMOTIONNEL", "SIGNAL DE DOULEUR",
    "RÉPONSE MINIMISANTE", "DÉTOUR CONVERSATIONNEL", "EFFET MIROIR MÉCANIQUE",
    "1.1 , Réponse longue", "1.2 , Réponse courte", "1.3 , Réponse hors sujet",
    "Sous-cas A", "Sous-cas B", "Sous-cas C",
    "2.1 , Principe", "2.2 , Les deux formes de rondeur et quand les utiliser",
    "2.3 , Quand mettre la rondeur avant la question, quand le mettre après",
    "3.1 , La règle fondamentale", "3.2 , Signaux d'engagement émotionnel et niveaux",
    "4.1 , Ce qu'est un signal de douleur",
    "4.2 , L'action quand un signal de douleur est détecté",
    "4.3 , Nuance : douleur verbalisée ≠ automatiquement bon client",
    "5.1 , Ce qu'est une réponse minimisante", "5.2 , L'action face à une réponse minimisante",
    "6.1 , Définition et déclencheur",
    "7.1 , Ce qu'il faut éviter", "7.2 , Ce qui fonctionne à la place",
    "INSIGHT PROFOND", "QUESTION À L'ENJEU PHILOSOPHIQUE", "LECTURE SUPERFICIELLE",
    "LECTURE PROFONDE", "FORMULATION EN OBSERVATION",
    "Source 1 , Les posts (source principale)", "Source 2 , Section info / bio",
    "Source 3 , Commentaires laissés sur d'autres posts",
    "Source 4 , Ancienneté et niveau d'activité",
    "✗ Ce qu'on ne fait PAS", "✓ Ce qu'on fait",
    "Niveau 1 , Lecture superficielle (à éviter)",
    "Niveau 2 , Lecture profonde (standard minimum requis)",
    "CAS B , Outbound (contact à froid)", "CAS C , Demande de connexion entrante",
    "CAS D , Visiteur de profil",
    "CAS A , Échange en commentaire préalable ⚠ EXCEPTION",
    "QUESTION OUVERTE", "QUESTION FERMÉE", "QUESTION DE COLLECTE",
    "QUESTION DE CONSCIENTISATION", "UNITÉ DE QUESTION",
    "Questions de collecte", "Questions de conscientisation",
    "Comment savoir quand passer au sujet suivant",
    "Ce qu'il ne faut pas faire", "Ce qui fonctionne à la place",
    "NATUREL SONORE", "REBOND D'ÉNERGIE", "RONDEUR MAL PLACÉE",
    "PERSONNALISATION SUPERFICIELLE", "IMPERFECTIONS VOLONTAIRES",
    "La règle",
    "Rondeur mal placée, ce qu'il ne faut pas faire", "Rondeur bien placée, le principe",
    "Personnalisation superficielle", "Validation excessive", "Pitch déguisé",
    "Fausse sincérité", "Sujets financiers directs", "Patterns IA",
    "Ponctuation", "Emojis", "Imperfections volontaires", "Sauts de ligne",
    "Test 1, Le test à l'oral", "Test 2, Le test de la rondeur",
    "Test 3, Le test de la personnalisation", "Test 4 , Le test IA",
    "BESOIN RÉEL", "SOLVABILITÉ", "SPECTRE DES ANGLES D'ATTAQUE", "LEAD SCORING", "CLARIFICATION",
    "1.1 , Ce qui confirme un besoin réel", "1.2 , Deux niveaux de conscientisation",
    "1.3 , Le spectre des angles d'attaque", "2.1 , Principe",
    "3.1 , Critères minimum avant de proposer l'action finale",
    "3.2 , Trop tôt / trop tard", "3.3 , Formulations urgence",
    "Famille 1 , Projection temporelle", "Famille 2 , Situation actuelle",
    "Famille 3 , Coût de l'inaction", "Famille 4 , Ancrée dans les mots du prospect",
    "4.1 , La temporalité n'est pas un problème indépendant",
    "4.2 , Les trois issues après clarification",
    "4.3 , Six cas types",
    "4.4 , 'Non' mou vs 'non' ferme, même réflexe de départ",
    "4.5 , Suivi des prospects 'pas le bon timing'",
    "Issue 1 , Vraie contrainte confirmée", "Issue 2 , Objection de surface révélée",
    "Issue 3 , Prospect reste vague ou résiste",
    "CAS 1 , Budget alloué ailleurs", "CAS 2 , Pivot d'offre en cours",
    "CAS 3 , Bande passante personnelle saturée",
    "CAS 4 , Déjà engagé dans un autre accompagnement",
    "CAS 5 , Période creuse saisonnière", "CAS 6 , Gros client en cours",
    "Étape 1 , Message à la date convenue",
    "Étape 2 , Si pas de réponse après [X jours]", "Étape 3 , Si toujours pas de réponse",
    "DISQUALIFICATION", "ESPACE DE VÉRIFICATION", "SIGNAL D'ALERTE", "RELANCE",
    "1.1 , Trois relances sans réponse",
    "1.2 , Entonnoir inversé : clarification confirmée",
    "1.2b , Entonnoir classique : spectre épuisé",
    "1.3 , Fermeture agressive",
    "SÉQUENCE DE RELANCES", "Relance 1 :", "Relance 2 :", "Relance 3 :",
    "2.1 , Profil hors niche découvert en cours de conversation",
    "2.2 , Plaintes répétées sur les anciens coaches",
    "2.3 , Questions de garantie répétées",
    "2.4 , Prospect en mode collecte d'informations",
    "FERMETURE AGRESSIVE", "REFUS POLI", "SITUATION COUVERTE", "PATTERN DÉTECTÉ",
    "DÉSENGAGEMENT PROGRESSIF", "DÉSALIGNEMENT EXPRIMÉ",
    # Special cases from Bloc 4A examples
    "Exemple 1, Elsa", "Exemple 2, Danaë",
    # Majuscules and sub headings
    "Majuscules et minuscules", "Longueur adaptive",
}

# Paragraphs that should be NORMAL + ITALIC (Niveau 3 and numbered examples 1-9)
NORMAL_ITALIC_TITLES = {
    "Niveau 3 : Message ultra-personnalisé par insight profond",
}

paragraphs = doc.paragraphs
h1_color = (0x1F, 0x38, 0x64)
h2_color = (0x2E, 0x40, 0x57)
h3_color = (0x40, 0x40, 0x40)

def para_stripped(p):
    return p.text.strip().rstrip('\n').strip()

applied = {'h1': 0, 'h2': 0, 'h3': 0}

for p in paragraphs:
    txt = para_stripped(p)
    # Remove newlines for matching
    txt_clean = txt.replace('\n', ' ').strip()

    # H1 check
    matched_h1 = False
    for title in H1_TITLES:
        if txt_clean == title or (title in txt_clean and len(txt_clean) < len(title) + 5):
            try:
                p.style = doc.styles['Heading 1']
                apply_heading_style(p, 1, h1_color, 16, 24, 12)
                applied['h1'] += 1
                matched_h1 = True
            except Exception as e:
                pass
            break
    if matched_h1:
        continue

    # H2 check
    matched_h2 = False
    for title in H2_TITLES:
        if title.startswith("VOCABULAIRE DU BLOC"):
            if txt_clean.startswith("VOCABULAIRE DU BLOC"):
                try:
                    p.style = doc.styles['Heading 2']
                    apply_heading_style(p, 2, h2_color, 13, 18, 6)
                    applied['h2'] += 1
                    matched_h2 = True
                except Exception as e:
                    pass
                break
        elif txt_clean == title:
            try:
                p.style = doc.styles['Heading 2']
                apply_heading_style(p, 2, h2_color, 13, 18, 6)
                applied['h2'] += 1
                matched_h2 = True
            except Exception as e:
                pass
            break
    if matched_h2:
        continue

    # H3 check
    for title in H3_TITLES:
        if txt_clean == title:
            try:
                p.style = doc.styles['Heading 3']
                apply_heading_style(p, 3, h3_color, 11, 12, 4)
                applied['h3'] += 1
            except Exception as e:
                pass
            break

    # Niveau 3 italic check
    if txt_clean in NORMAL_ITALIC_TITLES:
        try:
            p.style = doc.styles['Normal']
        except KeyError:
            pass
        for run in p.runs:
            run.italic = True

print(f"  -> H1: {applied['h1']}, H2: {applied['h2']}, H3: {applied['h3']}")

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 6 : FORMATAGE DES EXEMPLES (italique + indentation)
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 6 : formatage des exemples...")

# Patterns that indicate an example message follows
EXAMPLE_TRIGGERS = {'template :', 'script :', 'exemple :', 'formulation :'}

# Patterns that look like direct message text (conversational, lowercase starts)
MESSAGE_PATTERNS = [
    r'^salut \[',
    r'^salut [a-záàâéèêëîïôùûüç]',
    r'^ouais ',
    r'^je me demande',
    r'^c\'est normal',
    r'^ok, je comprends',
    r'^j\'espère que',
    r'^zéro souci',
    r'^\[prénom\]',
    r'^hey ',
]

paragraphs = doc.paragraphs
n_examples = 0
in_example_block = False
example_count = 0

for i, p in enumerate(paragraphs):
    txt = p.text.strip().lower()
    txt_orig = p.text.strip()

    # Check if this is a trigger line
    is_trigger = any(txt.startswith(t) for t in EXAMPLE_TRIGGERS)

    if is_trigger:
        # Make trigger bold
        for run in p.runs:
            run.bold = True
        in_example_block = True
        example_count = 0
        continue

    # Check if this looks like a message example
    is_message = False
    if in_example_block and example_count < 3 and txt_orig:
        is_message = True
        example_count += 1

    # Also check by pattern
    for pat in MESSAGE_PATTERNS:
        if re.match(pat, txt_orig, re.IGNORECASE):
            is_message = True
            break

    if is_message and txt_orig:
        # Apply italic to all runs
        for run in p.runs:
            run.italic = True
        # Apply left indent
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:left'), '284')  # ~0.5cm in twips (1cm=567 twips, 0.5cm≈284)
        n_examples += 1
    elif not txt_orig:
        # Empty line resets example block after 2 empties
        pass
    else:
        in_example_block = False
        example_count = 0

print(f"  -> {n_examples} exemples formatés")

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 7 : FORMATAGE IF/THEN/ELSE
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 7 : formatage IF/THEN/ELSE...")

paragraphs = doc.paragraphs
n_ifthen = 0

for p in paragraphs:
    txt = p.text
    if not txt:
        continue

    stripped = txt.lstrip()
    starts_with_if = stripped.startswith('IF ')
    starts_with_then = stripped.startswith('THEN ')
    starts_with_else = stripped.startswith('ELSE ')

    if starts_with_if or starts_with_then or starts_with_else:
        # Apply indentation
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:left'), '284')

        # Bold/underline the keyword in runs
        for run in p.runs:
            rt = run.text
            if rt.lstrip().startswith('IF ') and starts_with_if:
                run.bold = True
                n_ifthen += 1
            elif rt.lstrip().startswith('THEN ') and starts_with_then:
                run.bold = True
                run.underline = True
                n_ifthen += 1
            elif rt.lstrip().startswith('ELSE ') and starts_with_else:
                run.bold = True
                n_ifthen += 1

print(f"  -> {n_ifthen} blocs IF/THEN/ELSE formatés")

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 8 : SAUVEGARDE TEMPORAIRE
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 8 : sauvegarde intermédiaire...")
doc.save(TMP)
print("  -> Sauvegardé")

# ──────────────────────────────────────────────────────────────────────────────
# ÉTAPE 9 : GESTION DES COMMENTAIRES WORD (manipulation XML directe)
# ──────────────────────────────────────────────────────────────────────────────
print("Étape 9 : gestion des commentaires Word...")

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def w(tag):
    return f'{{{W}}}{tag}'

# New comments content
COMMENTS = [
    {
        'id': '1',
        'author': 'Matthias',
        'date': '2026-03-09T10:00:00Z',
        'text': 'À retravailler avec Kevin, A/B testing prévu. La formulation actuelle est trop générique, elle manque de personnalisation par rapport au profil du prospect.',
        'anchor_text': 'Tu vas être focus sur quoi dans les prochains mois ?',
        'label': 'A',
    },
    {
        'id': '2',
        'author': 'Matthias',
        'date': '2026-03-09T10:01:00Z',
        'text': 'Seuils de solvabilité précis (investissement minimum selon l\'offre) à définir avec Kevin en config plateforme. Les formulations concrètes pour qualifier la solvabilité indirectement varient selon la niche et l\'offre, à affiner lors d\'un atelier séparé.',
        'anchor_text': 'Seuils de solvabilité précis',
        'label': 'B',
    },
    {
        'id': '3',
        'author': 'Matthias',
        'date': '2026-03-09T10:02:00Z',
        'text': 'Lead scoring précis (critères pondérés, seuils exacts) à affiner avec Kevin. Pas encore défini.',
        'anchor_text': 'Lead scoring précis',
        'label': 'C',
    },
    {
        'id': '4',
        'author': 'Matthias',
        'date': '2026-03-09T10:03:00Z',
        'text': 'La phase post-clarification (cas types et scripts) est à travailler avec Kevin lors d\'un atelier séparé. Inclura de l\'A/B testing sur les formulations.',
        'anchor_text': 'Phase après la clarification',
        'label': 'D',
    },
    {
        'id': '5',
        'author': 'Matthias',
        'date': '2026-03-09T10:04:00Z',
        'text': 'Cadence exacte (nombre de jours entre chaque étape, nombre de tentatives selon le cas de timing) à paramétrer en config plateforme avec Kevin et Tariel. Le tronc central pose la logique, pas les paramètres techniques.',
        'anchor_text': 'La cadence entre l\'étape 1 et l\'étape 2 est configurable',
        'label': 'E',
    },
    {
        'id': '6',
        'author': 'Matthias',
        'date': '2026-03-09T10:05:00Z',
        'text': 'Le spectre des angles est propre à chaque niche et offre. Défini en config utilisateur ou dans les modules niche, pas ici.',
        'anchor_text': 'Le spectre exact est défini par l\'utilisateur',
        'label': 'F',
    },
    {
        'id': '7',
        'author': 'Matthias',
        'date': '2026-03-09T10:06:00Z',
        'text': 'À définir avec Kevin : quels facteurs déclenchent une remontée, comment l\'interface présente l\'info, et comment éviter que l\'utilisateur se retrouve avec 20 vérifications par jour.',
        'anchor_text': 'Facteurs qui déclenchent une remontée',
        'label': 'G',
    },
    {
        'id': '8',
        'author': 'Matthias',
        'date': '2026-03-09T10:07:00Z',
        'text': 'Cadence des relances à paramétrer en config plateforme avec Kevin (nombre de jours entre chaque envoi selon le profil et le stade).',
        'anchor_text': 'Cadence des relances',
        'label': 'H',
    },
    {
        'id': '9',
        'author': 'Matthias',
        'date': '2026-03-09T10:08:00Z',
        'text': 'À définir avec Kevin : la séquence de clarification et creusage avant d\'arriver à la disqualification, pour s\'assurer qu\'on a vraiment épuisé les angles.',
        'anchor_text': 'Séquence de clarification',
        'label': 'I',
    },
    {
        'id': '10',
        'author': 'Matthias',
        'date': '2026-03-09T10:09:00Z',
        'text': 'Angles concrets dans les modules niche et la config utilisateur. Clarifier aussi qui décide si un prospect "pourrait devenir pertinent plus tard" : toujours remonter dans l\'espace de vérification.',
        'anchor_text': 'Ce critère dépend entièrement de la configuration utilisateur',
        'label': 'J',
    },
    {
        'id': '11',
        'author': 'Matthias',
        'date': '2026-03-09T10:10:00Z',
        'text': '(1) Vocabulaire de qualification : certains mots déclenchent de la méfiance selon la niche. Liste à construire lors d\'un atelier dédié. (2) Mécanique de suivi long terme pour les prospects CAS 4/5/6 : voir Section 4.5. (3) Lead scoring précis : voir commentaire Section 3.2.',
        'anchor_text': '(1) Vocabulaire de qualification',
        'label': 'K',
    },
    {
        'id': '12',
        'author': 'Matthias',
        'date': '2026-03-09T10:11:00Z',
        'text': '(1) Formulations carte sur table : Section 3, situation Pattern détecté. (2) Formulations fermetures agressives : Sections 1.3 et 3. (3) Espace de vérification utilisateur : voir commentaire Vocabulaire Bloc 7. (4) Mécanique des 3 relances : Section 1.1, cadence à paramétrer avec Kevin. (5) Récap global des formulations : à produire en fin de tous les blocs sous forme de document de synthèse.',
        'anchor_text': 'Formulations carte sur table : traitées dans Section 3',
        'label': 'L',
    },
]

def make_comment_xml(comment_id, author, date, text):
    """Create a w:comment XML element."""
    c = etree.Element(w('comment'))
    c.set(w('id'), str(comment_id))
    c.set(w('author'), author)
    c.set(w('date'), date)
    c.set(w('initials'), author[0])

    p_elem = etree.SubElement(c, w('p'))
    pPr = etree.SubElement(p_elem, w('pPr'))
    pStyle = etree.SubElement(pPr, w('pStyle'))
    pStyle.set(w('val'), 'CommentText')

    r_elem = etree.SubElement(p_elem, w('r'))
    rPr = etree.SubElement(r_elem, w('rPr'))
    rStyle = etree.SubElement(rPr, w('rStyle'))
    rStyle.set(w('val'), 'CommentReference')
    t_ref = etree.SubElement(r_elem, w('t'))
    t_ref.text = str(comment_id)

    r_text = etree.SubElement(p_elem, w('r'))
    t_text = etree.SubElement(r_text, w('t'))
    t_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_text.text = text

    return c

# Work with the zip file
with zipfile.ZipFile(TMP, 'r') as z:
    names = z.namelist()
    file_contents = {}
    for name in names:
        file_contents[name] = z.read(name)

# Parse document.xml
doc_xml = etree.fromstring(file_contents['word/document.xml'])

# Remove existing comment markup from document.xml
for tag in ['commentRangeStart', 'commentRangeEnd', 'commentReference']:
    for elem in doc_xml.findall(f'.//{w(tag)}'):
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)

# Find paragraphs in document XML and add comment anchors
body = doc_xml.find(f'.//{w("body")}')
all_doc_paras = doc_xml.findall(f'.//{w("p")}')

def para_text_from_xml(p_elem):
    texts = []
    for t in p_elem.findall(f'.//{w("t")}'):
        if t.text:
            texts.append(t.text)
    return ''.join(texts)

def add_comment_to_para(para_elem, comment_id):
    """Add commentRangeStart/End and commentReference to a paragraph."""
    # Add commentRangeStart at beginning of paragraph
    crs = etree.Element(w('commentRangeStart'))
    crs.set(w('id'), str(comment_id))

    cre = etree.Element(w('commentRangeEnd'))
    cre.set(w('id'), str(comment_id))

    # Create run with comment reference
    r_ref = etree.Element(w('r'))
    rPr = etree.SubElement(r_ref, w('rPr'))
    rStyle = etree.SubElement(rPr, w('rStyle'))
    rStyle.set(w('val'), 'CommentReference')
    cr = etree.SubElement(r_ref, w('commentReference'))
    cr.set(w('id'), str(comment_id))

    # Insert at appropriate positions
    # Find first run in paragraph
    runs = para_elem.findall(w('r'))
    pPr = para_elem.find(w('pPr'))

    # Insert commentRangeStart after pPr
    if pPr is not None:
        pPr_idx = list(para_elem).index(pPr)
        para_elem.insert(pPr_idx + 1, crs)
    else:
        para_elem.insert(0, crs)

    # Append commentRangeEnd and reference at end
    para_elem.append(cre)
    para_elem.append(r_ref)

# Match comments to paragraphs
comments_placed = {}
for comment in COMMENTS:
    anchor_text = comment['anchor_text']
    comment_id = comment['id']

    # Find best matching paragraph
    best_para = None
    best_score = 0

    for para_elem in all_doc_paras:
        ptxt = para_text_from_xml(para_elem)
        if anchor_text in ptxt:
            score = len(anchor_text)
            if score > best_score:
                best_score = score
                best_para = para_elem

    if best_para is not None:
        add_comment_to_para(best_para, comment_id)
        comments_placed[comment['label']] = True
        print(f"  Comment {comment['label']} placé sur: {para_text_from_xml(best_para)[:60]}")
    else:
        print(f"  ⚠ Comment {comment['label']} non placé (texte anchor non trouvé: '{anchor_text[:50]}')")

# Build new comments.xml
nsmap = {
    'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
    'mo': 'http://schemas.microsoft.com/office/mac/office/2008/main',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'mv': 'urn:schemas-microsoft-com:mac:vml',
    'o': 'urn:schemas-microsoft-com:office:office',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'v': 'urn:schemas-microsoft-com:vml',
    'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'w10': 'urn:schemas-microsoft-com:office:word',
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
    'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
    'wpi': 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk',
    'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
}
comments_root = etree.Element(w('comments'), nsmap=nsmap)

for comment in COMMENTS:
    c_elem = make_comment_xml(comment['id'], comment['author'], comment['date'], comment['text'])
    comments_root.append(c_elem)

# Serialize back
new_doc_xml = etree.tostring(doc_xml, xml_declaration=True, encoding='UTF-8', standalone=True)
new_comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding='UTF-8', standalone=True)

# Write new zip
import io
output_buf = io.BytesIO()
with zipfile.ZipFile(output_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
    for name in names:
        if name == 'word/document.xml':
            zout.writestr(name, new_doc_xml)
        elif name == 'word/comments.xml':
            zout.writestr(name, new_comments_xml)
        else:
            zout.writestr(name, file_contents[name])
    # If comments.xml didn't exist, add it
    if 'word/comments.xml' not in names:
        zout.writestr('word/comments.xml', new_comments_xml)

output_buf.seek(0)
with open(DST, 'wb') as f:
    f.write(output_buf.read())

import os
size = os.path.getsize(DST)
print(f"\n✅ Fichier généré : {DST}")
print(f"   Taille : {size:,} octets ({size/1024:.1f} KB)")
print(f"   Commentaires placés : {len(comments_placed)}/12")
